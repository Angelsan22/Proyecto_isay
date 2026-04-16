from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse, Response
from sqlalchemy.orm import Session
from sqlalchemy import func
import io
import os
from datetime import datetime
from typing import Optional

from app import models, database
from app.database import get_db

router = APIRouter(
    prefix="/reportes",
    tags=["Administradores"]
)
def get_report_data(tipo: str, db: Session):
    if tipo == "ventas":
        pedidos = db.query(models.Pedido).all()
        data = [
            {"ID": p.id, "Fecha": str(p.fecha)[:10], "Cliente ID": p.cliente_id, "Total": p.total, "Estatus": p.estatus} 
            for p in pedidos
        ]
        columns = ["ID", "Fecha", "Cliente ID", "Total", "Estatus"]
        title = "Reporte General de Ventas"
        
    elif tipo == "inventario":
        productos = db.query(models.Producto).all()
        data = [
            {"ID": p.id, "Nombre": p.nombre, "Stock Actual": p.stock_actual, "Stock Minimo": p.stock_minimo, "Categoria": p.categoria}
            for p in productos
        ]
        columns = ["ID", "Nombre", "Stock Actual", "Stock Minimo", "Categoria"]
        title = "Reporte de Inventario / Stock"

    elif tipo == "clientes":
        res = db.query(
            models.Usuario.nombre, 
            func.count(models.Pedido.id).label("total_pedidos"),
            func.sum(models.Pedido.total).label("total_gastado")
        ).join(models.Pedido).group_by(models.Usuario.id).order_by(func.sum(models.Pedido.total).desc()).all()
        
        data = [{"Nombre": r[0], "Pedidos": r[1], "Total Gastado": r[2]} for r in res]
        columns = ["Nombre", "Pedidos", "Total Gastado"]
        title = "Ranking de Clientes Top"

    elif tipo == "estatus":
        res = db.query(models.Pedido.estatus, func.count(models.Pedido.id)).group_by(models.Pedido.estatus).all()
        data = [{"Estatus": r[0], "Cantidad": r[1]} for r in res]
        columns = ["Estatus", "Cantidad"]
        title = "Distribución de Pedidos por Estatus"
    
    else:
        return None, None, None

    return data, columns, title

def generate_pdf(data, columns, title):
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, title, ln=True, align="C")
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 10, f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True, align="R")
    pdf.ln(10)
    pdf.set_fill_color(232, 103, 27)
    pdf.set_text_color(255, 255, 255)
    col_width = 190 / len(columns)
    for col in columns:
        pdf.cell(col_width, 10, col, border=1, fill=True)
    pdf.ln()
    pdf.set_text_color(0, 0, 0)
    for row in data:
        for col in columns:
            pdf.cell(col_width, 8, str(row.get(col, "")), border=1)
        pdf.ln()
    
    return pdf.output(dest='S')

def generate_xlsx(data, columns, title):
    import pandas as pd
    df = pd.DataFrame(data, columns=columns)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Reporte')
    return output.getvalue()

def generate_docx(data, columns, title):
    from docx import Document
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    table = doc.add_table(rows=1, cols=len(columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
    
    for row_data in data:
        row_cells = table.add_row().cells
        for i, col in enumerate(columns):
            row_cells[i].text = str(row_data.get(col, ""))
    
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

@router.get("/descargar/{tipo}")
def descargar_reporte(
    tipo: str, 
    formato: str = Query("pdf", enum=["pdf", "xlsx", "docx"]),
    db: Session = Depends(get_db)
):
    data, columns, title = get_report_data(tipo, db)
    if not data:
        raise HTTPException(status_code=404, detail="Tipo de reporte no válido")

    filename = f"reporte_{tipo}_{datetime.now().strftime('%Y%m%d')}.{formato}"

    if formato == "pdf":
        content = generate_pdf(data, columns, title)
        if isinstance(content, bytearray):
            content = bytes(content)
        media_type = "application/pdf"
    elif formato == "xlsx":
        content = generate_xlsx(data, columns, title)
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif formato == "docx":
        content = generate_docx(data, columns, title)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
